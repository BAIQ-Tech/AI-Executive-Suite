"""
Document Processing Service

Handles document upload, processing, and context extraction
for the AI Executive Suite.
"""

import logging
import os
import tempfile
from typing import List, Dict, Optional, Any, BinaryIO, Union
from dataclasses import dataclass
from datetime import datetime
from enum import Enum
import hashlib
import mimetypes
import uuid
from pathlib import Path

# Document processing libraries
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
import openpyxl
from openpyxl import load_workbook
import csv
import io

# File type detection
HAS_MAGIC = False
magic = None
try:
    import magic
    HAS_MAGIC = True
except ImportError:
    # python-magic not available, will use fallback methods
    pass

# Security scanning
import re

logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """Document type classifications"""
    FINANCIAL = "financial"
    TECHNICAL = "technical"
    STRATEGIC = "strategic"
    LEGAL = "legal"
    OPERATIONAL = "operational"
    OTHER = "other"


class SensitivityLevel(Enum):
    """Document sensitivity levels"""
    PUBLIC = "public"
    INTERNAL = "internal"
    CONFIDENTIAL = "confidential"
    RESTRICTED = "restricted"


@dataclass
class FileUpload:
    """File upload data"""
    filename: str
    content: Union[BinaryIO, bytes]
    content_type: str
    size: int
    
    def get_content_bytes(self) -> bytes:
        """Get content as bytes"""
        if isinstance(self.content, bytes):
            return self.content
        elif hasattr(self.content, 'read'):
            self.content.seek(0)
            data = self.content.read()
            self.content.seek(0)
            return data
        else:
            raise ValueError("Invalid content type")


@dataclass
class DocumentMetadata:
    """Document metadata"""
    title: Optional[str] = None
    description: Optional[str] = None
    tags: List[str] = None
    document_type: Optional[DocumentType] = None
    sensitivity_level: SensitivityLevel = SensitivityLevel.INTERNAL
    author: Optional[str] = None
    department: Optional[str] = None


@dataclass
class Document:
    """Document model"""
    id: str
    user_id: str
    filename: str
    file_type: str
    file_size: int
    content_hash: str
    extracted_text: str
    summary: str
    key_insights: List[str]
    document_type: DocumentType
    sensitivity_level: SensitivityLevel
    embedding_id: Optional[str]
    created_at: datetime
    processed_at: Optional[datetime]
    last_accessed: Optional[datetime]
    reference_count: int
    decisions_referenced: List[str]
    metadata: Dict[str, Any]


@dataclass
class DocumentContext:
    """Context extracted from document"""
    document_id: str
    content: str
    relevance_score: float
    page_number: Optional[int] = None
    section: Optional[str] = None


@dataclass
class DocumentFilters:
    """Filters for document search"""
    document_types: Optional[List[DocumentType]] = None
    sensitivity_levels: Optional[List[SensitivityLevel]] = None
    date_range: Optional[tuple] = None
    tags: Optional[List[str]] = None
    author: Optional[str] = None


@dataclass
class DocumentSearchResult:
    """Document search result"""
    document: Document
    relevance_score: float
    matching_excerpts: List[str]


class FileProcessingError(Exception):
    """Custom exception for file processing errors"""
    pass


class SecurityScanError(Exception):
    """Custom exception for security scan failures"""
    pass


class DocumentProcessingService:
    """Service for document processing and management"""
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.logger = logging.getLogger(__name__)
        self.upload_dir = config.get('upload_directory', 'uploads')
        self.max_file_size = config.get('max_file_size', 50 * 1024 * 1024)  # 50MB default
        self.allowed_extensions = config.get('allowed_extensions', [
            'pdf', 'docx', 'doc', 'xlsx', 'xls', 'txt', 'csv'
        ])
        
        # Log magic library availability
        if not HAS_MAGIC:
            self.logger.warning("python-magic not available, falling back to extension-based file type detection")
        
        # Ensure upload directory exists
        os.makedirs(self.upload_dir, exist_ok=True)
        
        self.supported_types = {
            'application/pdf': 'pdf',
            'application/msword': 'doc',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'application/vnd.ms-excel': 'xls',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
            'text/plain': 'txt',
            'text/csv': 'csv'
        }
        
        # Security patterns to detect potentially malicious content
        self.security_patterns = [
            re.compile(r'<script[^>]*>.*?</script>', re.IGNORECASE | re.DOTALL),
            re.compile(r'javascript:', re.IGNORECASE),
            re.compile(r'vbscript:', re.IGNORECASE),
            re.compile(r'on\w+\s*=', re.IGNORECASE),
            re.compile(r'<iframe[^>]*>.*?</iframe>', re.IGNORECASE | re.DOTALL),
            re.compile(r'<object[^>]*>.*?</object>', re.IGNORECASE | re.DOTALL),
            re.compile(r'<embed[^>]*>', re.IGNORECASE),
        ]
        
    def upload_document(
        self, 
        file_upload: FileUpload, 
        metadata: DocumentMetadata,
        user_id: str
    ) -> Document:
        """
        Upload and process a document with comprehensive validation and processing
        
        Args:
            file_upload: File upload data
            metadata: Document metadata
            user_id: ID of uploading user
            
        Returns:
            Document object with processing results
            
        Raises:
            FileProcessingError: If file processing fails
            SecurityScanError: If security scan fails
            ValueError: If validation fails
        """
        try:
            self.logger.info(f"Processing document upload: {file_upload.filename}")
            
            # Step 1: Validate file
            self._validate_file(file_upload)
            
            # Step 2: Detect actual file type (security measure)
            actual_file_type = self._detect_file_type(file_upload)
            
            # Step 3: Security scan
            self._security_scan(file_upload)
            
            # Step 4: Generate content hash for deduplication
            content_hash = self._generate_content_hash(file_upload.content)
            
            # Step 5: Save file to storage
            file_path = self._save_file(file_upload, content_hash)
            
            # Step 6: Extract text content
            extracted_text = self._extract_text(file_upload, actual_file_type)
            
            # Step 7: Generate summary and insights (placeholder for now)
            summary = self._generate_summary(extracted_text)
            key_insights = self._extract_key_insights(extracted_text)
            
            # Step 8: Classify document type if not provided
            document_type = metadata.document_type or self._classify_document(extracted_text)
            
            # Step 9: Create document object
            document = Document(
                id=str(uuid.uuid4()),
                user_id=user_id,
                filename=self._sanitize_filename(file_upload.filename),
                file_type=actual_file_type,
                file_size=file_upload.size,
                content_hash=content_hash,
                extracted_text=extracted_text,
                summary=summary,
                key_insights=key_insights,
                document_type=document_type,
                sensitivity_level=metadata.sensitivity_level,
                embedding_id=None,  # Will be set after vector processing
                created_at=datetime.now(),
                processed_at=datetime.now(),
                last_accessed=None,
                reference_count=0,
                decisions_referenced=[],
                metadata={
                    'title': metadata.title,
                    'description': metadata.description,
                    'tags': metadata.tags or [],
                    'author': metadata.author,
                    'department': metadata.department,
                    'file_path': file_path,
                    'original_filename': file_upload.filename
                }
            )
            
            # Step 10: Generate embeddings for semantic search
            self._generate_embeddings(document)
            
            self.logger.info(f"Successfully processed document: {document.id}")
            return document
            
        except Exception as e:
            self.logger.error(f"Error processing document {file_upload.filename}: {str(e)}")
            raise FileProcessingError(f"Failed to process document: {str(e)}")
    
    def extract_context(
        self, 
        document_id: str, 
        query: str,
        max_results: int = 5
    ) -> List[DocumentContext]:
        """
        Extract relevant context from a document based on query using vector search
        
        Args:
            document_id: ID of the document
            query: Search query or context request
            max_results: Maximum number of context pieces to return
            
        Returns:
            List of relevant document contexts
        """
        self.logger.info(f"Extracting context from document {document_id} for query: {query[:100]}")
        
        try:
            # Import here to avoid circular imports
            from services.vector_database import VectorDatabaseService
            
            # Get vector database configuration
            vector_config = {
                'openai_api_key': self.config.get('openai_api_key'),
                'embedding_model': self.config.get('embedding_model', 'text-embedding-3-small'),
                'chroma_path': self.config.get('chroma_path', './chroma_db'),
                'collection_name': self.config.get('collection_name', 'ai_executive_documents')
            }
            
            vector_service = VectorDatabaseService(vector_config)
            
            # Search for relevant context
            search_results = vector_service.get_document_context(
                document_id=document_id,
                query=query,
                max_chunks=max_results
            )
            
            # Convert to DocumentContext objects
            contexts = []
            for result in search_results:
                context = DocumentContext(
                    document_id=result.document_id,
                    content=result.content,
                    relevance_score=result.similarity_score,
                    page_number=None,  # Would need to be calculated from chunk position
                    section=f"Chunk {result.chunk_index}"
                )
                contexts.append(context)
            
            self.logger.info(f"Found {len(contexts)} relevant contexts for document {document_id}")
            return contexts
            
        except Exception as e:
            self.logger.warning(f"Vector search failed for document {document_id}: {str(e)}")
            # Fallback to placeholder
            return [
                DocumentContext(
                    document_id=document_id,
                    content="Context extraction temporarily unavailable",
                    relevance_score=0.5,
                    page_number=None,
                    section="Fallback"
                )
            ]
    
    def search_documents(
        self, 
        query: str, 
        filters: DocumentFilters = None,
        max_results: int = 10
    ) -> List[DocumentSearchResult]:
        """
        Search documents using semantic search
        
        Args:
            query: Search query
            filters: Optional filters to apply
            max_results: Maximum number of results
            
        Returns:
            List of matching documents with relevance scores
        """
        self.logger.info(f"Searching documents for query: {query[:100]}")
        
        try:
            # Import here to avoid circular imports
            from services.vector_database import VectorDatabaseService
            
            # Get vector database configuration
            vector_config = {
                'openai_api_key': self.config.get('openai_api_key'),
                'embedding_model': self.config.get('embedding_model', 'text-embedding-3-small'),
                'chroma_path': self.config.get('chroma_path', './chroma_db'),
                'collection_name': self.config.get('collection_name', 'ai_executive_documents')
            }
            
            vector_service = VectorDatabaseService(vector_config)
            
            # Build metadata filter
            metadata_filter = {}
            if filters:
                if filters.document_types:
                    metadata_filter['document_type'] = {
                        '$in': [dt.value for dt in filters.document_types]
                    }
                if filters.sensitivity_levels:
                    metadata_filter['sensitivity_level'] = {
                        '$in': [sl.value for sl in filters.sensitivity_levels]
                    }
            
            # Search for similar content
            search_results = vector_service.search_similar_content(
                query=query,
                n_results=max_results,
                metadata_filter=metadata_filter if metadata_filter else None
            )
            
            # Group results by document and create DocumentSearchResult objects
            document_results = {}
            for result in search_results:
                doc_id = result.document_id
                if doc_id not in document_results:
                    # Create a placeholder Document object (would normally fetch from database)
                    placeholder_doc = Document(
                        id=doc_id,
                        user_id="",
                        filename=result.metadata.get('filename', 'Unknown'),
                        file_type=result.metadata.get('file_type', 'unknown'),
                        file_size=0,
                        content_hash="",
                        extracted_text="",
                        summary="",
                        key_insights=[],
                        document_type=DocumentType.OTHER,
                        sensitivity_level=SensitivityLevel.INTERNAL,
                        embedding_id="",
                        created_at=datetime.now(),
                        processed_at=None,
                        last_accessed=None,
                        reference_count=0,
                        decisions_referenced=[],
                        metadata=result.metadata
                    )
                    
                    document_results[doc_id] = DocumentSearchResult(
                        document=placeholder_doc,
                        relevance_score=result.similarity_score,
                        matching_excerpts=[result.content]
                    )
                else:
                    # Add excerpt and update relevance score
                    document_results[doc_id].matching_excerpts.append(result.content)
                    # Use highest relevance score
                    document_results[doc_id].relevance_score = max(
                        document_results[doc_id].relevance_score,
                        result.similarity_score
                    )
            
            results = list(document_results.values())
            # Sort by relevance score
            results.sort(key=lambda x: x.relevance_score, reverse=True)
            
            self.logger.info(f"Found {len(results)} documents matching query")
            return results
            
        except Exception as e:
            self.logger.warning(f"Vector search failed: {str(e)}")
            # Return empty results on failure
            return []
    
    def get_document_by_id(self, document_id: str) -> Optional[Document]:
        """
        Get document by ID
        
        Args:
            document_id: ID of the document
            
        Returns:
            Document object or None if not found
        """
        self.logger.info(f"Fetching document {document_id}")
        
        # Placeholder implementation
        return None
    
    def delete_document(self, document_id: str, user_id: str) -> bool:
        """
        Delete a document
        
        Args:
            document_id: ID of the document
            user_id: ID of user requesting deletion
            
        Returns:
            True if successful, False otherwise
        """
        self.logger.info(f"Deleting document {document_id}")
        
        # Placeholder implementation
        return True
    
    def _validate_file(self, file_upload: FileUpload) -> None:
        """
        Validate uploaded file for size, type, and basic security
        
        Args:
            file_upload: File upload data
            
        Raises:
            ValueError: If validation fails
        """
        # Check file size
        if file_upload.size > self.max_file_size:
            raise ValueError(f"File size {file_upload.size} exceeds maximum allowed size {self.max_file_size}")
        
        if file_upload.size == 0:
            raise ValueError("File is empty")
        
        # Check filename
        if not file_upload.filename or len(file_upload.filename.strip()) == 0:
            raise ValueError("Filename is required")
        
        # Check file extension
        file_ext = Path(file_upload.filename).suffix.lower().lstrip('.')
        if file_ext not in self.allowed_extensions:
            raise ValueError(f"File extension '{file_ext}' not allowed. Allowed: {', '.join(self.allowed_extensions)}")
        
        # Basic filename security check
        dangerous_chars = ['..', '/', '\\', '<', '>', ':', '"', '|', '?', '*']
        if any(char in file_upload.filename for char in dangerous_chars):
            raise ValueError("Filename contains dangerous characters")
    
    def _detect_file_type(self, file_upload: FileUpload) -> str:
        """
        Detect actual file type using magic numbers (more secure than relying on MIME type)
        
        Args:
            file_upload: File upload data
            
        Returns:
            Detected file type
            
        Raises:
            ValueError: If file type cannot be determined or is not supported
        """
        try:
            content_bytes = file_upload.get_content_bytes()
            
            if HAS_MAGIC:
                # Use python-magic to detect file type
                mime_type = magic.from_buffer(content_bytes, mime=True)
                
                # Map MIME type to our supported types
                if mime_type in self.supported_types:
                    return self.supported_types[mime_type]
            else:
                # Fallback: basic signature detection
                mime_type = self._detect_mime_by_signature(content_bytes)
                if mime_type and mime_type in self.supported_types:
                    return self.supported_types[mime_type]
            
            # Fallback to extension-based detection
            file_ext = Path(file_upload.filename).suffix.lower().lstrip('.')
            if file_ext in self.allowed_extensions:
                return file_ext
            
            raise ValueError(f"Unsupported file type: {mime_type if 'mime_type' in locals() else 'unknown'}")
            
        except Exception as e:
            self.logger.error(f"Error detecting file type: {str(e)}")
            raise ValueError(f"Could not determine file type: {str(e)}")
    
    def _detect_mime_by_signature(self, content_bytes: bytes) -> Optional[str]:
        """
        Basic MIME type detection using file signatures (fallback when python-magic is not available)
        
        Args:
            content_bytes: File content as bytes
            
        Returns:
            Detected MIME type or None
        """
        if len(content_bytes) < 4:
            return None
        
        # Common file signatures
        signatures = {
            b'\x25\x50\x44\x46': 'application/pdf',  # PDF
            b'\x50\x4B\x03\x04': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',  # DOCX/XLSX
            b'\xD0\xCF\x11\xE0': 'application/msword',  # DOC/XLS
        }
        
        for signature, mime_type in signatures.items():
            if content_bytes.startswith(signature):
                return mime_type
        
        # Check if it's plain text
        try:
            content_bytes[:1024].decode('utf-8')
            return 'text/plain'
        except UnicodeDecodeError:
            pass
        
        return None
    
    def _security_scan(self, file_upload: FileUpload) -> None:
        """
        Perform security scan on uploaded file
        
        Args:
            file_upload: File upload data
            
        Raises:
            SecurityScanError: If security threats are detected
        """
        try:
            content_bytes = file_upload.get_content_bytes()
            
            # Check for executable file signatures
            executable_signatures = [
                b'\x4d\x5a',  # PE executable
                b'\x7f\x45\x4c\x46',  # ELF executable
                b'\xca\xfe\xba\xbe',  # Mach-O executable
                b'\xfe\xed\xfa\xce',  # Mach-O executable (reverse)
            ]
            
            for signature in executable_signatures:
                if content_bytes.startswith(signature):
                    raise SecurityScanError("Executable files are not allowed")
            
            # Convert to text for pattern scanning (if possible)
            try:
                text_content = content_bytes.decode('utf-8', errors='ignore')
                
                # Check for malicious patterns
                for pattern in self.security_patterns:
                    if pattern.search(text_content):
                        raise SecurityScanError(f"Potentially malicious content detected: {pattern.pattern}")
                        
            except UnicodeDecodeError:
                # Binary file, skip text-based security checks
                pass
            
            self.logger.info(f"Security scan passed for file: {file_upload.filename}")
            
        except SecurityScanError:
            raise
        except Exception as e:
            self.logger.error(f"Error during security scan: {str(e)}")
            raise SecurityScanError(f"Security scan failed: {str(e)}")
    
    def _sanitize_filename(self, filename: str) -> str:
        """
        Sanitize filename for safe storage
        
        Args:
            filename: Original filename
            
        Returns:
            Sanitized filename
        """
        # Remove dangerous characters
        sanitized = re.sub(r'[<>:"/\\|?*]', '_', filename)
        
        # Remove leading/trailing whitespace and dots
        sanitized = sanitized.strip(' .')
        
        # Ensure filename is not empty
        if not sanitized:
            sanitized = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        return sanitized
    
    def _save_file(self, file_upload: FileUpload, content_hash: str) -> str:
        """
        Save uploaded file to storage
        
        Args:
            file_upload: File upload data
            content_hash: SHA-256 hash of file content
            
        Returns:
            File path where the file was saved
        """
        # Create subdirectory based on hash prefix for better organization
        subdir = os.path.join(self.upload_dir, content_hash[:2])
        os.makedirs(subdir, exist_ok=True)
        
        # Generate unique filename using hash and original extension
        file_ext = Path(file_upload.filename).suffix
        filename = f"{content_hash}{file_ext}"
        file_path = os.path.join(subdir, filename)
        
        # Save file
        content_bytes = file_upload.get_content_bytes()
        with open(file_path, 'wb') as f:
            f.write(content_bytes)
        
        self.logger.info(f"File saved to: {file_path}")
        return file_path
    
    def _generate_content_hash(self, content: Union[BinaryIO, bytes]) -> str:
        """
        Generate SHA-256 hash of file content
        
        Args:
            content: File content as BinaryIO or bytes
            
        Returns:
            SHA-256 hash as hex string
        """
        hash_obj = hashlib.sha256()
        
        if isinstance(content, bytes):
            hash_obj.update(content)
        else:
            content.seek(0)
            for chunk in iter(lambda: content.read(4096), b""):
                hash_obj.update(chunk)
            content.seek(0)
        
        return hash_obj.hexdigest()
    
    def _extract_text(self, file_upload: FileUpload, file_type: str) -> str:
        """
        Extract text from uploaded file based on file type
        
        Args:
            file_upload: File upload data
            file_type: Detected file type
            
        Returns:
            Extracted text content
            
        Raises:
            FileProcessingError: If text extraction fails
        """
        try:
            content_bytes = file_upload.get_content_bytes()
            
            if file_type == 'pdf':
                return self._extract_pdf_text(content_bytes)
            elif file_type == 'docx':
                return self._extract_docx_text(content_bytes)
            elif file_type == 'doc':
                return self._extract_doc_text(content_bytes)
            elif file_type in ['xlsx', 'xls']:
                return self._extract_excel_text(content_bytes)
            elif file_type == 'txt':
                return self._extract_txt_text(content_bytes)
            elif file_type == 'csv':
                return self._extract_csv_text(content_bytes)
            else:
                raise FileProcessingError(f"Text extraction not implemented for file type: {file_type}")
                
        except Exception as e:
            self.logger.error(f"Error extracting text from {file_type} file: {str(e)}")
            raise FileProcessingError(f"Failed to extract text: {str(e)}")
    
    def _extract_pdf_text(self, content_bytes: bytes) -> str:
        """Extract text from PDF file"""
        text_parts = []
        
        try:
            # Try with pdfplumber first (better for complex layouts)
            with io.BytesIO(content_bytes) as pdf_file:
                with pdfplumber.open(pdf_file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
        except Exception as e:
            self.logger.warning(f"pdfplumber failed, trying PyPDF2: {str(e)}")
            
            # Fallback to PyPDF2
            try:
                with io.BytesIO(content_bytes) as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
            except Exception as e2:
                raise FileProcessingError(f"Failed to extract PDF text with both libraries: {str(e2)}")
        
        return '\n\n'.join(text_parts)
    
    def _extract_docx_text(self, content_bytes: bytes) -> str:
        """Extract text from DOCX file"""
        text_parts = []
        
        with io.BytesIO(content_bytes) as docx_file:
            doc = DocxDocument(docx_file)
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
        
        return '\n\n'.join(text_parts)
    
    def _extract_doc_text(self, content_bytes: bytes) -> str:
        """Extract text from DOC file (legacy Word format)"""
        # Note: python-docx doesn't support .doc files
        # For production, you might want to use python-docx2txt or antiword
        raise FileProcessingError("Legacy .doc format not supported. Please convert to .docx format.")
    
    def _extract_excel_text(self, content_bytes: bytes) -> str:
        """Extract text from Excel file"""
        text_parts = []
        
        with io.BytesIO(content_bytes) as excel_file:
            workbook = load_workbook(excel_file, read_only=True, data_only=True)
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"Sheet: {sheet_name}")
                
                # Extract data from cells
                for row in sheet.iter_rows(values_only=True):
                    row_data = []
                    for cell_value in row:
                        if cell_value is not None:
                            row_data.append(str(cell_value))
                    if row_data:
                        text_parts.append(' | '.join(row_data))
                
                text_parts.append('')  # Empty line between sheets
        
        return '\n'.join(text_parts)
    
    def _extract_txt_text(self, content_bytes: bytes) -> str:
        """Extract text from plain text file"""
        try:
            # Try UTF-8 first
            return content_bytes.decode('utf-8')
        except UnicodeDecodeError:
            try:
                # Fallback to latin-1
                return content_bytes.decode('latin-1')
            except UnicodeDecodeError:
                # Last resort: ignore errors
                return content_bytes.decode('utf-8', errors='ignore')
    
    def _extract_csv_text(self, content_bytes: bytes) -> str:
        """Extract text from CSV file"""
        text_parts = []
        
        try:
            # Try UTF-8 first
            text_content = content_bytes.decode('utf-8')
        except UnicodeDecodeError:
            text_content = content_bytes.decode('latin-1')
        
        csv_reader = csv.reader(io.StringIO(text_content))
        for row in csv_reader:
            if row:  # Skip empty rows
                text_parts.append(' | '.join(str(cell) for cell in row))
        
        return '\n'.join(text_parts)
    
    def _generate_summary(self, text: str) -> str:
        """
        Generate document summary using advanced analysis
        
        Args:
            text: Extracted text content
            
        Returns:
            Document summary
        """
        try:
            # Import here to avoid circular imports
            from services.document_analysis import DocumentAnalysisService, AnalysisType
            
            # Get analysis configuration
            analysis_config = {
                'openai_api_key': self.config.get('openai_api_key'),
                'analysis_model': self.config.get('analysis_model', 'gpt-3.5-turbo')
            }
            
            analysis_service = DocumentAnalysisService(analysis_config)
            
            # Generate summary
            summary_result = analysis_service.generate_summary(text)
            return summary_result.detailed_summary
            
        except Exception as e:
            self.logger.warning(f"Advanced summary generation failed: {str(e)}")
            # Fallback to simple summary
            if len(text) <= 500:
                return text
            
            # Find a good breaking point near 500 characters
            summary = text[:500]
            last_sentence = summary.rfind('.')
            if last_sentence > 300:  # If we found a sentence end reasonably close
                summary = summary[:last_sentence + 1]
            
            return summary + "..."
    
    def _extract_key_insights(self, text: str) -> List[str]:
        """
        Extract key insights from document using advanced analysis
        
        Args:
            text: Extracted text content
            
        Returns:
            List of key insights
        """
        try:
            # Import here to avoid circular imports
            from services.document_analysis import DocumentAnalysisService, AnalysisType
            
            # Get analysis configuration
            analysis_config = {
                'openai_api_key': self.config.get('openai_api_key'),
                'analysis_model': self.config.get('analysis_model', 'gpt-3.5-turbo')
            }
            
            analysis_service = DocumentAnalysisService(analysis_config)
            
            # Extract insights
            insights_result = analysis_service.extract_key_insights(text)
            
            # Convert to string list
            insights = [insight.insight for insight in insights_result]
            
            # Add document statistics
            word_count = len(text.split())
            insights.append(f"Document contains approximately {word_count:,} words")
            
            return insights if insights else ["Document processed successfully"]
            
        except Exception as e:
            self.logger.warning(f"Advanced insight extraction failed: {str(e)}")
            # Fallback to simple keyword-based extraction
            insights = []
            
            # Simple keyword-based insight extraction
            financial_keywords = ['revenue', 'profit', 'cost', 'budget', 'investment', 'roi', 'financial']
            technical_keywords = ['system', 'architecture', 'technology', 'development', 'software', 'infrastructure']
            strategic_keywords = ['strategy', 'market', 'competition', 'growth', 'opportunity', 'risk']
            
            text_lower = text.lower()
            
            if any(keyword in text_lower for keyword in financial_keywords):
                insights.append("Contains financial information and metrics")
            
            if any(keyword in text_lower for keyword in technical_keywords):
                insights.append("Includes technical specifications and system details")
            
            if any(keyword in text_lower for keyword in strategic_keywords):
                insights.append("Discusses strategic planning and market analysis")
            
            # Add word count insight
            word_count = len(text.split())
            insights.append(f"Document contains approximately {word_count:,} words")
            
            return insights if insights else ["Document processed successfully"]
    
    def _classify_document(self, text: str) -> DocumentType:
        """
        Classify document type based on content using advanced analysis
        
        Args:
            text: Extracted text content
            
        Returns:
            Classified document type
        """
        try:
            # Import here to avoid circular imports
            from services.document_analysis import DocumentAnalysisService, AnalysisType
            
            # Get analysis configuration
            analysis_config = {
                'openai_api_key': self.config.get('openai_api_key'),
                'analysis_model': self.config.get('analysis_model', 'gpt-3.5-turbo')
            }
            
            analysis_service = DocumentAnalysisService(analysis_config)
            
            # Categorize document
            category_result = analysis_service.categorize_document(text)
            
            # Map analysis categories to DocumentType enum
            category_mapping = {
                'financial': DocumentType.FINANCIAL,
                'technical': DocumentType.TECHNICAL,
                'strategic': DocumentType.STRATEGIC,
                'legal': DocumentType.LEGAL,
                'operational': DocumentType.OPERATIONAL,
                'marketing': DocumentType.OTHER  # No marketing type in enum
            }
            
            return category_mapping.get(category_result.primary_category, DocumentType.OTHER)
            
        except Exception as e:
            self.logger.warning(f"Advanced document classification failed: {str(e)}")
            # Fallback to simple keyword-based classification
            text_lower = text.lower()
            
            # Simple keyword-based classification
            financial_keywords = ['financial', 'budget', 'revenue', 'profit', 'cost', 'investment', 'balance sheet', 'income statement']
            technical_keywords = ['technical', 'system', 'architecture', 'software', 'development', 'api', 'database']
            strategic_keywords = ['strategy', 'strategic', 'market', 'business plan', 'roadmap', 'vision', 'mission']
            legal_keywords = ['contract', 'agreement', 'legal', 'terms', 'conditions', 'compliance', 'regulation']
            operational_keywords = ['process', 'procedure', 'operations', 'workflow', 'manual', 'guide']
            
            # Count keyword matches
            financial_score = sum(1 for keyword in financial_keywords if keyword in text_lower)
            technical_score = sum(1 for keyword in technical_keywords if keyword in text_lower)
            strategic_score = sum(1 for keyword in strategic_keywords if keyword in text_lower)
            legal_score = sum(1 for keyword in legal_keywords if keyword in text_lower)
            operational_score = sum(1 for keyword in operational_keywords if keyword in text_lower)
            
            # Return type with highest score
            scores = {
                DocumentType.FINANCIAL: financial_score,
                DocumentType.TECHNICAL: technical_score,
                DocumentType.STRATEGIC: strategic_score,
                DocumentType.LEGAL: legal_score,
                DocumentType.OPERATIONAL: operational_score
            }
            
            max_score = max(scores.values())
            if max_score > 0:
                return max(scores, key=scores.get)
            
            return DocumentType.OTHER
    
    def _generate_embeddings(self, document: Document) -> str:
        """Generate vector embeddings for document using vector database service"""
        try:
            # Import here to avoid circular imports
            from services.vector_database import VectorDatabaseService
            
            # Get vector database configuration
            vector_config = {
                'openai_api_key': self.config.get('openai_api_key'),
                'embedding_model': self.config.get('embedding_model', 'text-embedding-3-small'),
                'chroma_path': self.config.get('chroma_path', './chroma_db'),
                'collection_name': self.config.get('collection_name', 'ai_executive_documents')
            }
            
            vector_service = VectorDatabaseService(vector_config)
            
            # Create embeddings for the document
            chunk_ids = vector_service.create_document_embeddings(
                document_id=document.id,
                content=document.extracted_text,
                metadata={
                    'filename': document.filename,
                    'file_type': document.file_type,
                    'document_type': document.document_type.value if document.document_type else None,
                    'sensitivity_level': document.sensitivity_level.value if document.sensitivity_level else None,
                    'created_at': document.created_at.isoformat() if document.created_at else None
                }
            )
            
            # Set embedding ID to the first chunk ID (for reference)
            embedding_id = f"vector_chunks_{len(chunk_ids)}"
            document.embedding_id = embedding_id
            
            self.logger.info(f"Generated {len(chunk_ids)} vector embeddings for document {document.id}")
            return embedding_id
            
        except Exception as e:
            self.logger.warning(f"Failed to generate embeddings for document {document.id}: {str(e)}")
            # Set a placeholder embedding ID
            embedding_id = f"embedding_failed_{document.id}"
            document.embedding_id = embedding_id
            return embedding_id